import streamlit as st
import fitz
import re
import io
import hashlib
import numpy as np
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="시험지 짜깁기", layout="wide")
st.title("시험지 문제 선택기")
st.caption("PDF 시험지에서 원하는 문제만 골라 새 시험지를 만들어 드립니다.")

# ── 사이드바 ───────────────────────────────────────────────────
NUM_PATTERN = r"^(\d{1,2})[\.\。．\)\） ]"
X_LIMIT_PCT = 30

with st.sidebar:
    st.header("레이아웃 설정")
    two_col = st.checkbox("2단 레이아웃 (좌/우 컬럼)", value=True)

    st.divider()
    st.header("출력 설정")
    cover_on   = st.checkbox("표지 제작", value=False)
    cover_text = ""
    if cover_on:
        cover_text = st.text_area("표지 문구",
            placeholder="예: 2024학년도 1학기\n수학 시험",
            height=120)
    separate_sources = st.checkbox("파일 간 구분 (페이지 단위)", value=False,
        help="여러 PDF를 넣었을 때 파일이 바뀌는 지점에서 새 페이지로 분리합니다.")
    divider_line = st.checkbox("세로 구분선", value=False,
        help="좌/우 컬럼 사이 가운데에 세로선을 그립니다.")
    show_page_num = st.checkbox("쪽 번호 표시", value=False,
        help="각 페이지 하단 가운데에 쪽 번호를 표시합니다.")
    st.subheader("페이지 여백 (cm)")
    wm_top    = st.number_input("위",     0.0, 5.0, 1.5, step=0.1)
    wm_bottom = st.number_input("아래",   0.0, 5.0, 1.5, step=0.1)
    wm_left   = st.number_input("왼쪽",  0.0, 5.0, 1.5, step=0.1)
    wm_right  = st.number_input("오른쪽", 0.0, 5.0, 1.5, step=0.1)

    st.divider()
    debug_mode = st.checkbox("진단 모드", value=False)

# ── PDF 업로드 ─────────────────────────────────────────────────
uploaded_files = st.file_uploader(
    "PDF 시험지를 업로드하세요 (여러 개 가능)",
    type=["pdf"],
    accept_multiple_files=True,
)
if not uploaded_files:
    st.stop()


# ── 유틸: 픽셀 기반 tight clip ────────────────────────────────
def tight_clip(page, cx0: float, cx1: float, y_lo: float, y_hi: float,
               pad_y: float = 3.0):
    h_pt = y_hi - y_lo
    w_pt = cx1 - cx0
    if h_pt <= 0 or w_pt <= 0:
        return fitz.Rect(cx0, y_lo, cx1, y_hi)

    pix = page.get_pixmap(clip=fitz.Rect(cx0, y_lo, cx1, y_hi), colorspace=fitz.csGRAY)
    H, W = pix.height, pix.width
    if H == 0 or W == 0:
        return fitz.Rect(cx0, y_lo, cx1, y_hi)

    arr  = np.frombuffer(pix.samples, dtype=np.uint8).reshape(H, W)
    dark = arr < 200

    nw_col   = dark.sum(axis=0)
    col_mask = (nw_col > 0) & (nw_col / H < 0.6)
    c_idx    = np.where(col_mask)[0]
    if len(c_idx) == 0:
        return fitz.Rect(cx0, y_lo, cx1, y_hi)
    c0, c1 = int(c_idx[0]), int(c_idx[-1]) + 1

    inner  = dark[:, c0:c1]
    iW     = inner.shape[1]
    nw_row = inner.sum(axis=1)
    r_mask = (nw_row > 2) & (nw_row / iW < 0.88)
    r_idx  = np.where(r_mask)[0]
    if len(r_idx) == 0:
        return fitz.Rect(cx0, y_lo, cx1, y_hi)

    excl_mask = (nw_row > 2) & (nw_row / iW >= 0.88)
    excl_idx  = np.where(excl_mask)[0]
    bot_sep   = excl_idx[excl_idx > H * 0.80] if len(excl_idx) > 0 else np.array([], dtype=int)
    ABOVE_WIN = 25
    for sep_cand in bot_sep:
        sep_row    = int(sep_cand)
        near_above = r_idx[(r_idx >= sep_row - ABOVE_WIN) & (r_idx < sep_row)]
        if len(near_above) == 0:
            r_idx = r_idx[r_idx < sep_row]
            if len(r_idx) == 0:
                return fitz.Rect(cx0, y_lo, cx1, y_hi)
            break

    r_first = int(r_idx[0])
    r_last  = int(r_idx[0])
    for i in range(len(r_idx) - 1):
        if int(r_idx[i + 1]) - int(r_idx[i]) > 70:
            break
        r_last = int(r_idx[i + 1])

    c_dark = dark[r_first:r_last + 1, :]
    cH     = c_dark.shape[0]
    if cH > 0:
        nw_c2  = c_dark.sum(axis=0)
        keep   = (nw_c2 > 0) & (nw_c2 / cH < 0.80)
        c2_idx = np.where(keep)[0]
        if len(c2_idx) > 0:
            c0 = int(c2_idx[0])
            c1 = int(c2_idx[-1]) + 1

    y_sc = h_pt / H
    x_sc = w_pt / W
    ty0  = max(y_lo, y_lo + r_first * y_sc - pad_y)
    ty1  = min(y_hi, y_lo + (r_last + 1) * y_sc + pad_y)
    tx0  = max(cx0,  cx0  + c0 * x_sc)
    tx1  = min(cx1,  cx0  + c1 * x_sc)
    return fitz.Rect(tx0, ty0, tx1, ty1)


# ── 문제 위치 감지 ─────────────────────────────────────────────
@st.cache_data(show_spinner="문제 번호 감지 중...", max_entries=20)
def find_problems(data: bytes, pdf_hash: str, pattern: str, x_pct: float, two_col: bool) -> dict:
    doc  = fitz.open(stream=data, filetype="pdf")
    locs = []
    seen = set()

    for pi in range(len(doc)):
        page = doc[pi]
        pw   = page.rect.width
        half = pw / 2
        thr  = pw * x_pct / 100

        for block in page.get_text("dict")["blocks"]:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                x0       = line["bbox"][0]
                in_left  = x0 < thr
                in_right = two_col and (half < x0 < half + thr)
                if not (in_left or in_right):
                    continue
                txt = "".join(s["text"] for s in line["spans"]).strip()
                m   = re.match(pattern, txt)
                if not m:
                    continue
                try:
                    n = int(m.group(1))
                except (ValueError, IndexError):
                    continue
                if not (1 <= n <= 150) or n in seen:
                    continue
                seen.add(n)
                col_x0 = half if in_right else 0.0
                col_x1 = pw   if in_right else (half if two_col else pw)
                locs.append((pi, line["bbox"][1], n, col_x0, col_x1))

    if not locs:
        return {}

    locs.sort(key=lambda x: (x[0], x[1]))

    problems = {}
    for i, (pi, y0, n, cx0, cx1) in enumerate(locs):
        ph  = float(doc[pi].rect.height)
        ny0 = ph
        for j in range(i + 1, len(locs)):
            nxt_pi, nxt_y, _, nxt_cx0, _ = locs[j]
            if nxt_pi != pi:
                break
            if nxt_cx0 == cx0:
                ny0 = nxt_y
                break
        problems[n] = (pi, y0, ny0, cx0, cx1)

    return problems


# ── 번호 파싱 ──────────────────────────────────────────────────
def parse_nums(s: str, available: set) -> list:
    result = set()
    for part in s.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                result.update(range(int(a.strip()), int(b.strip()) + 1))
            except ValueError:
                pass
        elif part.isdigit():
            result.add(int(part))
    return sorted(result & available)


# ── 파일별 UI ─────────────────────────────────────────────────
# all_sources: list of (pdf_bytes, problems, selected)
all_sources = []

for fi, uf in enumerate(uploaded_files):
    pdf_bytes = uf.read()
    pdf_hash  = hashlib.md5(pdf_bytes).hexdigest()

    with st.expander(f"📄 {uf.name}", expanded=True):
        # 진단 모드
        if debug_mode:
            with st.expander("진단: 숫자로 시작하는 줄 목록", expanded=True):
                _doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                rows = []
                for _pi in range(len(_doc)):
                    _pw   = _doc[_pi].rect.width
                    _half = _pw / 2
                    _thr  = _pw * X_LIMIT_PCT / 100
                    for _blk in _doc[_pi].get_text("dict")["blocks"]:
                        if _blk["type"] != 0:
                            continue
                        for _ln in _blk["lines"]:
                            _txt = "".join(s["text"] for s in _ln["spans"]).strip()
                            if not _txt or not _txt[0].isdigit():
                                continue
                            _x0       = _ln["bbox"][0]
                            _in_left  = _x0 < _thr
                            _in_right = two_col and (_half < _x0 < _half + _thr)
                            _matched  = bool(re.match(NUM_PATTERN, _txt))
                            rows.append({
                                "페이지": _pi + 1,
                                "줄 텍스트": _txt[:50],
                                "x좌표": round(_x0, 1),
                                "y좌표": round(_ln["bbox"][1], 1),
                                "컬럼": "우" if _x0 >= _half else "좌",
                                "최종감지": "✅" if (_matched and (_in_left or _in_right)) else "❌",
                            })
                if rows:
                    st.dataframe(rows, use_container_width=True)
                else:
                    st.warning("숫자로 시작하는 줄이 없습니다.")

        problems = find_problems(pdf_bytes, pdf_hash, NUM_PATTERN, float(X_LIMIT_PCT), two_col)
        nums     = sorted(problems)

        if not problems:
            st.error("문제 번호를 감지하지 못했습니다. 사이드바 설정을 조정하거나 직접 입력 탭을 이용하세요.")
            all_sources.append((pdf_bytes, {}, []))
            continue

        st.success(f"문제 감지 완료: {nums[0]}번 ~ {nums[-1]}번 ({len(nums)}개)")

        st.subheader("포함할 문제 선택")
        tab_auto, tab_manual = st.tabs(["자동 선택 (감지된 문제)", "직접 번호 입력"])

        with tab_auto:
            b1, b2, _ = st.columns([1, 1, 4])
            if b1.button("전체 선택", key=f"sel_all_{fi}"):
                for n in nums:
                    st.session_state[f"cb_{fi}_{n}"] = True
            if b2.button("전체 해제", key=f"desel_all_{fi}"):
                for n in nums:
                    st.session_state[f"cb_{fi}_{n}"] = False

            r1, r2 = st.columns([4, 1])
            rng = r1.text_input("범위 추가", placeholder="예: 1,3,5-10,15",
                                label_visibility="collapsed", key=f"rng_{fi}")
            if r2.button("추가", key=f"add_rng_{fi}"):
                for part in rng.split(","):
                    part = part.strip()
                    if not part:
                        continue
                    if "-" in part:
                        try:
                            a, b = part.split("-", 1)
                            for n in range(int(a.strip()), int(b.strip()) + 1):
                                if n in problems:
                                    st.session_state[f"cb_{fi}_{n}"] = True
                        except ValueError:
                            pass
                    elif part.isdigit():
                        n = int(part)
                        if n in problems:
                            st.session_state[f"cb_{fi}_{n}"] = True

            grid = [nums[i:i + 10] for i in range(0, len(nums), 10)]
            for row in grid:
                cols = st.columns(10)
                for j, n in enumerate(row):
                    cols[j].checkbox(str(n), key=f"cb_{fi}_{n}")

            selected_auto = sorted(n for n in nums if st.session_state.get(f"cb_{fi}_{n}", False))
            if selected_auto:
                st.write(f"선택된 문제: **{selected_auto}**")
            else:
                st.info("체크박스에서 문제를 선택하세요.")

        with tab_manual:
            st.markdown("감지 여부와 상관없이 직접 번호를 입력합니다.")
            manual_input = st.text_input("문제 번호 입력",
                                         placeholder="예: 1,3,5-10,15",
                                         label_visibility="collapsed",
                                         key=f"manual_{fi}")
            selected_manual = parse_nums(manual_input, set(problems)) if manual_input else []
            if selected_manual:
                st.write(f"입력된 문제: **{selected_manual}**")

        final_selected = sorted(set(selected_auto + selected_manual))
        if final_selected:
            st.caption(f"이 파일에서 선택: {final_selected} ({len(final_selected)}개)")

        all_sources.append((pdf_bytes, problems, final_selected))


# ── 표지 생성 헬퍼 ─────────────────────────────────────────────
def _pdf_cover(out, text, A4_W, A4_H):
    import os
    page     = out.new_page(width=A4_W, height=A4_H)
    fontsize = 40
    lines    = text.strip().split("\n")
    line_h   = fontsize * 2.2
    total_h  = len(lines) * line_h
    y0       = (A4_H - total_h) / 2
    font_file = next(
        (p for p in ["C:/Windows/Fonts/malgunbd.ttf", "C:/Windows/Fonts/malgun.ttf"]
         if os.path.exists(p)), None)
    for i, line in enumerate(lines):
        rect = fitz.Rect(60, y0 + i * line_h, A4_W - 60, y0 + (i + 1) * line_h)
        kw   = dict(fontsize=fontsize, align=fitz.TEXT_ALIGN_CENTER, color=(0, 0, 0))
        if font_file:
            kw["fontfile"] = font_file
            kw["fontname"] = "coverfont"
        page.insert_textbox(rect, line, **kw)


# ── Word 헬퍼 ──────────────────────────────────────────────────
def remove_table_borders(table, divider_line=False):
    tbl        = table._tbl
    tblPr      = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH"]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    iv = OxmlElement("w:insideV")
    if divider_line:
        iv.set(qn("w:val"), "single")
        iv.set(qn("w:sz"), "4")
        iv.set(qn("w:color"), "AAAAAA")
    else:
        iv.set(qn("w:val"), "none")
    tblBorders.append(iv)
    tblPr.append(tblBorders)


# ── PDF 생성 ───────────────────────────────────────────────────
def make_pdf(sources, cover_text="", separate_sources=False, divider_line=False, show_page_num=False):
    """sources: list of (pdf_bytes, problems, selected)"""
    out = fitz.open()

    A4_W, A4_H = 595.0, 842.0
    pm_x, pm_y = 36.0, 36.0
    col_gap     = 20.0
    row_gap     = 10.0
    NROW, NCOL  = 2, 2
    slots_per   = NROW * NCOL

    col_w  = (A4_W - 2 * pm_x - (NCOL - 1) * col_gap) / NCOL
    area_h = (A4_H - 2 * pm_y - (NROW - 1) * row_gap) / NROW

    def slot_pos(global_slot):
        page_slot = global_slot % slots_per
        col = page_slot // NROW
        row = page_slot % NROW
        x0  = pm_x + col * (col_w + col_gap)
        y0  = pm_y + row * (area_h + row_gap)
        return global_slot // slots_per, x0, y0, row

    if cover_text.strip():
        _pdf_cover(out, cover_text, A4_W, A4_H)
        out.new_page(width=A4_W, height=A4_H)

    pages   = {}
    skipped = set()
    cur     = 0

    active_sources = [(b, p, s) for b, p, s in sources if s]
    for src_idx, (pdf_bytes, problems, selected) in enumerate(active_sources):
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        for n in selected:
            if n not in problems:
                continue
            while cur in skipped:
                cur += 1

            pi, y0, y1, cx0, cx1 = problems[n]
            clip  = tight_clip(doc[pi], cx0, cx1, y0, y1)
            if clip.width <= 0 or clip.height <= 0:
                cur += 1
                continue

            scale    = col_w / clip.width
            placed_h = clip.height * scale

            page_idx, x0, sy0, row = slot_pos(cur)

            if placed_h > area_h * 0.80 and row == NROW - 1:
                cur += 1
                while cur in skipped:
                    cur += 1
                page_idx, x0, sy0, row = slot_pos(cur)

            if page_idx not in pages:
                pages[page_idx] = out.new_page(width=A4_W, height=A4_H)

            dest = fitz.Rect(x0, sy0, x0 + col_w, sy0 + placed_h)
            pages[page_idx].show_pdf_page(dest, doc, pi, clip=clip)

            if placed_h > area_h * 0.80 and row < NROW - 1:
                skipped.add(cur + 1)

            cur += 1

        # 파일 구분: 다음 소스가 있으면 다음 페이지부터 시작
        if separate_sources and src_idx < len(active_sources) - 1:
            remainder = cur % slots_per
            if remainder != 0:
                cur += slots_per - remainder

    tmp = io.BytesIO()
    out.save(tmp, garbage=4, deflate=True)

    if divider_line or show_page_num:
        out2 = fitz.open(stream=tmp.getvalue(), filetype="pdf")
        cx = A4_W / 2
        cover_pages = 2 if cover_text.strip() else 0
        for pg_idx, pg in enumerate(out2):
            if divider_line:
                shape = pg.new_shape()
                shape.draw_line(fitz.Point(cx, pm_y), fitz.Point(cx, A4_H - pm_y))
                shape.finish(color=(0.7, 0.7, 0.7), width=0.5)
                shape.commit()
            if show_page_num and pg_idx >= cover_pages:
                num = pg_idx - cover_pages + 1
                pg.insert_textbox(
                    fitz.Rect(0, A4_H - pm_y, A4_W, A4_H),
                    str(num),
                    fontsize=9,
                    color=(0.4, 0.4, 0.4),
                    align=fitz.TEXT_ALIGN_CENTER,
                )
        buf = io.BytesIO()
        out2.save(buf, garbage=4, deflate=True)
        return buf.getvalue()

    return tmp.getvalue()


# ── Word 생성 ──────────────────────────────────────────────────
def make_word_doc(sources, wm_top, wm_bottom, wm_left, wm_right,
                  cover_text="", separate_sources=False, divider_line=False,
                  show_page_num=False, dpi=150):
    """sources: list of (pdf_bytes, problems, selected)"""
    word = Document()

    sec = word.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(wm_top)
    sec.bottom_margin = Cm(wm_bottom)
    sec.left_margin   = Cm(wm_left)
    sec.right_margin  = Cm(wm_right)

    NROW, NCOL  = 2, 2
    slots_per   = NROW * NCOL
    col_gap_cm  = 0.5
    row_gap_cm  = 0.5
    col_w_cm    = (21 - wm_left - wm_right - (NCOL - 1) * col_gap_cm) / NCOL
    area_h_cm   = (29.7 - wm_top - wm_bottom - (NROW - 1) * row_gap_cm) / NROW
    mat         = fitz.Matrix(dpi / 72, dpi / 72)

    def add_page_break():
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        word.add_paragraph().add_run()._r.append(br)

    if cover_text.strip():
        for i, line in enumerate(cover_text.strip().split("\n")):
            p = word.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(300) if i == 0 else Pt(8)
            run = p.add_run(line)
            run.font.size = Pt(36)
            run.font.bold = True
            run.font.name = "맑은 고딕"
        add_page_break()
        add_page_break()

    def cur_slot_info(s):
        page_slot = s % slots_per
        return s // slots_per, page_slot // NROW, page_slot % NROW

    page_grids = []
    skipped    = set()
    cur        = 0

    active_sources = [(b, p, s) for b, p, s in sources if s]
    for src_idx, (pdf_bytes, problems, selected) in enumerate(active_sources):
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        for n in selected:
            if n not in problems:
                continue
            while cur in skipped:
                cur += 1

            pi, y0, y1, cx0, cx1 = problems[n]
            clip = tight_clip(doc[pi], cx0, cx1, y0, y1)
            if clip.width <= 0 or clip.height <= 0:
                cur += 1
                continue

            placed_h_cm = col_w_cm * clip.height / clip.width
            page_idx, col, row = cur_slot_info(cur)

            if placed_h_cm > area_h_cm * 0.80 and row == NROW - 1:
                cur += 1
                while cur in skipped:
                    cur += 1
                page_idx, col, row = cur_slot_info(cur)

            while len(page_grids) <= page_idx:
                page_grids.append([[None] * NCOL for _ in range(NROW)])

            page_grids[page_idx][row][col] = (pdf_bytes, pi, cx0, cx1, y0, y1)

            if placed_h_cm > area_h_cm * 0.80 and row < NROW - 1:
                skipped.add(cur + 1)

            cur += 1

        # 파일 구분: 다음 소스가 있으면 다음 페이지부터 시작
        if separate_sources and src_idx < len(active_sources) - 1:
            remainder = cur % slots_per
            if remainder != 0:
                cur += slots_per - remainder

    for pg_idx, grid in enumerate(page_grids):
        if pg_idx > 0:
            add_page_break()

        table = word.add_table(rows=NROW, cols=NCOL)
        remove_table_borders(table, divider_line=divider_line)

        for row in range(NROW):
            for col in range(NCOL):
                cell_data = grid[row][col]
                if cell_data is None:
                    continue
                pdf_bytes2, pi2, cx0, cx1, y0, y1 = cell_data
                doc2  = fitz.open(stream=pdf_bytes2, filetype="pdf")
                clip2 = tight_clip(doc2[pi2], cx0, cx1, y0, y1)
                pix   = doc2[pi2].get_pixmap(matrix=mat, clip=clip2)
                para  = table.cell(row, col).paragraphs[0]
                para.add_run().add_picture(
                    io.BytesIO(pix.tobytes("png")), width=Cm(col_w_cm))

    if show_page_num:
        footer = word.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for tag, text in [("begin", None), (None, "PAGE"), ("end", None)]:
            if tag:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), tag)
            else:
                el = OxmlElement("w:instrText")
                el.text = text
            run._r.append(el)

    buf = io.BytesIO()
    word.save(buf)
    return buf.getvalue()


# ── 생성 버튼 ──────────────────────────────────────────────────
st.divider()

total_selected = sum(len(sel) for _, _, sel in all_sources)

if total_selected > 0:
    st.write(f"전체 선택: **{total_selected}개 문제**")

    c1, c2 = st.columns(2)

    if c1.button("📄 PDF로 내보내기", type="primary", use_container_width=True):
        with st.spinner("PDF 생성 중..."):
            pdf_out = make_pdf(all_sources, cover_text=cover_text, separate_sources=separate_sources, divider_line=divider_line, show_page_num=show_page_num)
        st.success("완료!")
        st.download_button(
            label="PDF 다운로드",
            data=pdf_out,
            file_name="new_exam.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

    if c2.button("📝 Word로 내보내기", use_container_width=True):
        with st.spinner("Word 생성 중..."):
            word_bytes = make_word_doc(
                all_sources,
                wm_top=wm_top, wm_bottom=wm_bottom,
                wm_left=wm_left, wm_right=wm_right,
                cover_text=cover_text,
                separate_sources=separate_sources,
                divider_line=divider_line,
                show_page_num=show_page_num,
            )
        st.success("완료!")
        st.download_button(
            label="Word 다운로드 (.docx)",
            data=word_bytes,
            file_name="new_exam.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
else:
    st.info("위에서 문제를 선택한 뒤 생성 버튼을 누르세요.")
